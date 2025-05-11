import os
import json
import re
import webbrowser
from dotenv import load_dotenv
from typing import Dict, Any, Optional, Union
from PyPDF2 import PdfReader
from docx import Document
import exifread

# -----------------------
# Constantes
# -----------------------
DICT_PATTERNS = {
   "Correo electrónico": r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+',
   "Dirección IP": r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b',
   "UUID": r'\b[a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12}\b',
   "Token/Clave": r'(?i)(api|auth|token|key|clave)[^\s]{5,}'
}

STR_OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Output")

# -----------------------
# Funciones de utilidad
# -----------------------
def fDetectSensitiveData(dMetadata: Dict[str, str]) -> None:
   """
   Analiza los valores de metadatos para detectar información sensible
   como correos, IPs, tokens, UUIDs, etc.
   
   Args:
      dMetadata: Diccionario con los metadatos extraídos
   """
   print("\nINFO    - Buscando posibles datos sensibles...")
   
   # Verificar si hay metadatos para analizar
   if not dMetadata:
      print("INFO    - No hay metadatos para analizar.")
      return
      
   bFound = False
   
   # Recorrer cada par clave-valor en los metadatos
   for sKey, sValue in dMetadata.items():
      if not isinstance(sValue, str):
         continue
         
      # Buscar patrones en cada valor
      for sTipo, sRegex in DICT_PATTERNS.items():
         list_matches = re.findall(sRegex, sValue)
         for sMatch in list_matches:
            print(f"ALERTA  - Posible {sTipo}: {sMatch}")
            bFound = True
   
   if not bFound:
      print("INFO    - No se encontraron datos sensibles.")

def fDmsToDecimal(tuple_dms, sRef: str) -> Optional[float]:
   """
   Convierte coordenadas GPS en formato grados/minutos/segundos (DMS)
   a decimal.
   
   Args:
      tuple_dms: Tupla con valores DMS
      sRef: Referencia geográfica (N, S, E, W)
      
   Returns:
      float: Coordenada en formato decimal, o None si hay error
   """
   try:
      flD = float(tuple_dms[0].num) / float(tuple_dms[0].den)
      flM = float(tuple_dms[1].num) / float(tuple_dms[1].den)
      flS = float(tuple_dms[2].num) / float(tuple_dms[2].den)
      
      # Cálculo de coordenada decimal
      flResult = flD + (flM / 60.0) + (flS / 3600.0)
      
      # Ajustar según hemisferio
      if sRef in ['S', 'W']:
         flResult = -flResult
         
      return flResult
   except Exception as e:
      print(f"ERROR   - Error al convertir coordenadas DMS: {e}")
      return None

def fShowGpsInMap(dExifData: Dict[str, Any]) -> None:
   """
   Extrae las coordenadas GPS del diccionario EXIF y abre Google Maps si las encuentra.
   
   Args:
      dExifData: Diccionario con datos EXIF de la imagen
   """
   try:
      # Obtener datos de latitud y longitud
      oLat = dExifData.get("GPS GPSLatitude")
      oLatRef = dExifData.get("GPS GPSLatitudeRef")
      oLon = dExifData.get("GPS GPSLongitude")
      oLonRef = dExifData.get("GPS GPSLongitudeRef")

      # Verificar que existen todos los datos necesarios
      if all([oLat, oLon, oLatRef, oLonRef]):
         flLat = fDmsToDecimal(oLat.values, oLatRef.values)
         flLon = fDmsToDecimal(oLon.values, oLonRef.values)
         
         if flLat is not None and flLon is not None:
            sUrl = f"https://www.google.com/maps?q={flLat},{flLon}"
            print(f"\nINFO    - Coordenadas GPS detectadas: {flLat}, {flLon}")
            print("INFO    - Abriendo ubicación en el navegador...")
            webbrowser.open(sUrl)
         else:
            print("INFO    - Coordenadas GPS no válidas.")
      else:
         print("INFO    - No se encontraron coordenadas GPS completas.")
   except Exception as e:
      print(f"ERROR   - No se pudo mostrar la ubicación GPS: {e}")

# -----------------------
# Funciones de extracción
# -----------------------
def fExtractPdfMetadata(sFilePath: str) -> Dict[str, str]:
   """
   Extrae metadatos de un archivo PDF usando PyPDF2
   
   Args:
      sFilePath: Ruta al archivo PDF
      
   Returns:
      dict: Diccionario con los metadatos del PDF
   """
   dMetadata = {}
   
   try:
      # Abrir y leer el PDF
      oReader = PdfReader(sFilePath)
      dRawMetadata = oReader.metadata or {}
      
      print(f"\nINFO    - PDF Metadata: {sFilePath}")
      
      # Convertir todos los valores a string para consistencia
      for sKey, oValue in dRawMetadata.items():
         # Eliminar el prefijo "/" que PyPDF2 agrega a las claves
         sCleanKey = sKey.strip("/") if sKey.startswith("/") else sKey
         sValue = str(oValue)
         print(f"INFO    - {sCleanKey}: {sValue}")
         dMetadata[sCleanKey] = sValue
         
      # Extraer información adicional si está disponible
      if oReader.pages:
         iPageCount = len(oReader.pages)
         print(f"INFO    - Páginas: {iPageCount}")
         dMetadata["PageCount"] = str(iPageCount)
         
      return dMetadata
   except Exception as e:
      print(f"ERROR   - Error leyendo PDF: {e}")
      return {}

def fExtractDocxMetadata(sFilePath: str) -> Dict[str, str]:
   """
   Extrae metadatos de un documento Word (.docx) usando python-docx
   
   Args:
      sFilePath: Ruta al archivo DOCX
      
   Returns:
      dict: Diccionario con los metadatos del documento
   """
   dMetadata = {}
   
   try:
      # Abrir y leer el documento Word
      oDoc = Document(sFilePath)
      oProps = oDoc.core_properties
      
      print(f"\nINFO    - Word Metadata: {sFilePath}")
      
      # Extraer propiedades del documento
      for sAttr in dir(oProps):
         # Filtrar solo propiedades reales (no métodos ni variables privadas)
         if not sAttr.startswith("_") and not callable(getattr(oProps, sAttr)):
            oValue = getattr(oProps, sAttr)
            
            # Mostrar solo propiedades con valores
            if oValue is not None:
               sValue = str(oValue)
               print(f"INFO    - {sAttr}: {sValue}")
               dMetadata[sAttr] = sValue
      
      # Agregar información adicional si está disponible
      if oDoc.paragraphs:
         iParagraphCount = len(oDoc.paragraphs)
         print(f"INFO    - Párrafos: {iParagraphCount}")
         dMetadata["ParagraphCount"] = str(iParagraphCount)
         
      return dMetadata
   except Exception as e:
      print(f"ERROR   - Error leyendo Word: {e}")
      return {}


def fExtractImageMetadata(sFilePath: str) -> Dict[str, str]:
   """
   Extrae metadatos EXIF de una imagen JPEG usando exifread.
   
   Args:
      sFilePath: Ruta al archivo de imagen
      
   Returns:
      dict: Diccionario con los metadatos EXIF
   """
   dMetadata = {}
   
   try:
      print(f"\nINFO    - Imagen Metadata: {sFilePath}")
      
      with open(sFilePath, 'rb') as file_obj:
         # Procesar el archivo EXIF con detalles
         dTags = exifread.process_file(file_obj, details=True)

         if dTags:
            # Convertir todos los valores a string para consistencia
            for sTag in sorted(dTags.keys()):
               sValue = str(dTags[sTag])
               print(f"INFO    - {sTag}: {sValue}")
               dMetadata[sTag] = sValue

            # Intentar mostrar mapa si hay coordenadas GPS
            fShowGpsInMap(dTags)
         else:
            print("INFO    - No se encontraron metadatos EXIF.")
               
      return dMetadata
   except Exception as e:
      print(f"ERROR   - Error leyendo imagen: {e}")
      return {}


def fSaveMetadataJson(dMetadata: Dict[str, str], sFilePath: str) -> Optional[str]:
   """
   Guarda los metadatos en un archivo JSON
   
   Args:
      dMetadata: Diccionario con los metadatos
      sFilePath: Ruta del archivo original
      
   Returns:
      str: Ruta del archivo JSON creado, o None si hubo error
   """
   try:
      # Crear directorio de salida si no existe
      os.makedirs(STR_OUTPUT_DIR, exist_ok=True)
      
      # Crear nombre de archivo basado en el nombre original
      sBaseName = os.path.splitext(os.path.basename(sFilePath))[0]
      sJsonFile = os.path.join(STR_OUTPUT_DIR, f"{sBaseName}_metadata.json")
      
      # Guardar metadatos en formato JSON
      with open(sJsonFile, 'w', encoding='utf-8') as file_out:
         json.dump(dMetadata, file_out, indent=4, ensure_ascii=False)
         
      print(f"\nINFO    - Metadatos exportados a: {sJsonFile}")
      return sJsonFile
   except Exception as e:
      print(f"ERROR   - No se pudo guardar el archivo JSON: {e}")
      return None


# -----------------------
# Función principal
# -----------------------
def fAnalyzeFile(sFilePath: str, bExportToJson: bool = True) -> Dict[str, str]:
   """
   Controlador principal: selecciona el tipo de archivo y extrae sus metadatos.
   También permite exportarlos a JSON y buscar datos sensibles.
   
   Args:
      sFilePath: Ruta al archivo que se va a analizar
      bExportToJson: Indica si se debe exportar a JSON
      
   Returns:
      dict: Diccionario con los metadatos extraídos
   """
   # Verificar que el archivo existe
   if not os.path.isfile(sFilePath):
      print(f"ERROR   - El archivo no existe: {sFilePath}")
      return {}

   # Determinar el tipo de archivo por su extensión
   sExt = os.path.splitext(sFilePath)[1].lower()
   dMetadata = {}

   # Seleccionar el extractor adecuado según la extensión
   if sExt == ".pdf":
      dMetadata = fExtractPdfMetadata(sFilePath)
   elif sExt == ".docx":
      dMetadata = fExtractDocxMetadata(sFilePath)
   elif sExt in [".jpg", ".jpeg", ".png", ".tiff", ".tif"]:
      dMetadata = fExtractImageMetadata(sFilePath)
   else:
      print(f"ERROR   - Tipo de archivo no soportado: {sExt}")
      return {}

   # Verificar si se obtuvieron metadatos
   if not dMetadata:
      print("WARNING - No se extrajo ningún metadato.")
      return {}

   # Exportar metadatos a archivo JSON si se solicitó
   if bExportToJson:
      fSaveMetadataJson(dMetadata, sFilePath)

   # Buscar datos sensibles en los metadatos
   fDetectSensitiveData(dMetadata)
   
   return dMetadata

# -----------------------
# Ejecutar desde IDE
# -----------------------
if __name__ == "__main__":
   # Cargar el archivo .env
   load_dotenv()
   # Obtener la ruta desde el archivo .env
   # Ruta del archivo que queremos analizar
   sTEST_FILE_PATH = os.getenv('sTEST_FILE_PATH')
   
   print("="*50)
   print("EXTRACTOR DE METADATOS")
   print("="*50)
   print(f"INFO    - Archivo a analizar: {sTEST_FILE_PATH}")
   print("INFO    - Extrayendo metadatos...")
   
   # Ejecutar el análisis
   dResult = fAnalyzeFile(sTEST_FILE_PATH, bExportToJson=True)
   
   print("\nINFO    - Análisis completado.")
   print("="*50)